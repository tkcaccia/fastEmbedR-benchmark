#!/usr/bin/env python3
"""Insert the current fastEmbedR performance-engineering text into the manuscript."""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table


def find_paragraph(document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix!r}")


def replace_paragraph(document, prefix: str, text: str) -> None:
    paragraph = find_paragraph(document, prefix)
    paragraph.text = text


def insert_paragraph_before(document, anchor, text: str, style: str):
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def format_improvement_table(table) -> None:
    widths = [Inches(value) for value in (1.22, 0.92, 0.92, 0.72, 2.52)]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(int(width.twips)))
        grid.append(column)

    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(width.twips for width in widths)))
    table_width.set(qn("w:type"), "dxa")

    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.width = widths[column_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index in (0, 4)
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8.25)
                    run.font.bold = row_index == 0


def populate_improvement_table(document):
    source = None
    for table in document.tables:
        if table.cell(0, 0).text.strip() == "Validation check":
            source = table
            break
    if source is None:
        raise ValueError("Could not find the five-column validation table template")

    table_element = deepcopy(source._tbl)
    table = Table(table_element, document._body)

    rows = [
        (
            "Component",
            "Baseline",
            "Retained",
            "Change",
            "Acceptance evidence",
        ),
        (
            "CPU openTSNE embedding",
            "16.259 s",
            "14.233 s",
            "-12.5%",
            "Coordinate correlation 1.000; Procrustes RMSD approximately 0; overlap@15 1.000.",
        ),
        (
            "CPU openTSNE, 1 versus 4 threads",
            "46.988 s",
            "14.304 s",
            "3.29x",
            "Same fixed KNN, initialization, iteration schedule, and objective; deterministic CPU result retained.",
        ),
        (
            "Metal openTSNE embedding",
            "2.602 s",
            "2.545 s",
            "-2.2%",
            "Correlation 0.99909-0.99967; normalized Procrustes RMSD 0.0257-0.0426; overlap@15 0.8816-0.9196.",
        ),
        (
            "CUDA openTSNE embedding",
            "3.293 s",
            "2.534 s",
            "-23.1%",
            "Correlation 0.999691; normalized Procrustes RMSD 0.024861; overlap@15 0.903091.",
        ),
        (
            "CUDA openTSNE full call",
            "5.234 s",
            "4.534 s",
            "-13.4%",
            "Same KNN policy, PCA initialization, FFT grid, and optimizer schedule.",
        ),
        (
            "Metal float32 PCA, warm",
            "0.429 s",
            "0.104 s",
            "-75.8% (4.1x)",
            "Minimum canonical correlation 0.9980; maximum principal angle 3.65 degrees; score-distance Pearson 0.9977.",
        ),
    ]
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
    format_improvement_table(table)
    return table


def update_document(source: Path, output: Path) -> None:
    document = Document(source)

    replacements = {
        "Modern biological and computational datasets": (
            "Modern biological and computational datasets often require dimensionality reduction methods that can handle large sample sizes while preserving meaningful local structure. fastEmbedR is a KNN-first R package for UMAP and openTSNE-style t-SNE workflows with native CPU, Metal, and CUDA embedding backends where available. The current implementation includes package-native CPU HNSW and Apple Metal exact/IVF-Flat nearest-neighbour search, while CUDA builds link directly to FAISS GPU and RAPIDS cuVS without calling Python. Reusable KNN matrices from faissR or another source can still be supplied to umap_knn() and opentsne_knn(). fastEmbedR implements KNN normalization, UMAP graph construction, openTSNE-style affinity construction, rSVD-based PCA initialization, float32 sparse buffers, layout optimization, backend validation, timing metadata, and transformation workflows. We evaluate fastEmbedR on a heterogeneous benchmark panel spanning image data, cytometry, metabolomics, foundation-model image features, and single-cell data. The benchmark compares total elapsed runtime and embedding quality against established R implementations, including Rtsne, FIt-SNE, umap, and uwot. Across successful matched comparisons, fastEmbedR provides a practical accelerated workflow, with the largest gains obtained from CUDA execution on large datasets. A controlled performance-engineering protocol further shows that implementation-only changes can reduce CPU, Metal, and CUDA runtime without reducing k, perplexity, epochs, or optimization iterations (van der Maaten and Hinton, 2008; McInnes et al., 2018; Policar et al., 2024; Caccia, 2026)."
        ),
        "Here, we present fastEmbedR": (
            "Here, we present fastEmbedR, a KNN-first embedding engine designed to make accelerated t-SNE and UMAP workflows available directly from R (Caccia, 2026). The package separates nearest-neighbour search from embedding at the API and source levels. Users may pass precomputed KNN indices and distances directly, which makes the UMAP and openTSNE-style embedding interfaces independent of a particular search engine. For one-call matrix workflows, the current package selects native float32 HNSW on CPU, native exact or recall-tuned IVF-Flat on Apple Metal, and direct FAISS GPU exact or RAPIDS cuVS IVF-Flat search on CUDA. CUDA KNN results can remain device-resident through graph or affinity construction. The optional faissR package remains useful when a separately reusable KNN object is required, but it is not called by the current one-call native paths. fastEmbedR then performs UMAP graph construction, openTSNE-style affinity construction, native layout optimization, backend reporting, PCA initialization, and transformation workflows. We evaluate these implementations on image, cytometry, metabolomics, foundation-model image-feature, and single-cell transcriptomic datasets. Simulated matrices are used only in separate nearest-neighbour stress tests and are not treated as evidence for embedding quality (van der Maaten and Hinton, 2008; McInnes et al., 2018; Johnson et al., 2019; Douze et al., 2024; Caccia, 2026)."
        ),
        "fastEmbedR is designed as a KNN-first embedding engine": (
            "fastEmbedR is designed as a KNN-first embedding engine rather than as a thin wrapper around existing dimensionality-reduction packages. Its main software contribution is a package-native execution path that performs nearest-neighbor selection for one-call workflows, KNN normalization, UMAP graph or openTSNE-style affinity construction, initialization, stochastic optimization, backend validation, timing metadata, and out-of-sample transformation from R. CPU one-call embeddings use native float32 HNSW; Metal uses native exact or recall-tuned IVF-Flat; CUDA builds link directly to FAISS GPU exact search and the RAPIDS cuVS C API for IVF-Flat. The optional faissR package is used only when users explicitly create a reusable KNN object. The exported pca() function uses package-owned randomized-SVD orchestration on CPU and CUDA and a resident float32 MPS block-subspace TSVD on Metal; the one-call CUDA openTSNE initializer separately uses native RAPIDS RAFT TSVD to retain scores on-device. No PCA route delegates to another R package, and none of the public CPU, Metal, or CUDA embedding paths call Python or reticulate (Johnson et al., 2019; Douze et al., 2024)."
        ),
        "For matrix input, backend selection is intentionally small.": (
            "For matrix input, backend selection is intentionally small. The public embedding backend argument is limited to backend = \"cpu\", \"metal\", or \"cuda\". CPU one-call embeddings use package-native HNSW. Metal uses package-native exact KNN below 4,096 observations and recall-tuned IVF-Flat for larger inputs. CUDA uses direct FAISS GPU exact search below 100,000 observations and recall-tuned RAPIDS cuVS IVF-Flat for larger inputs, with target_recall = 0.99. The CUDA result remains in device memory through graph or affinity construction and optimization. Users who require a reusable KNN graph or explicit search-method control call faissR::nn() and pass that object to umap_knn() or opentsne_knn()."
        ),
        "The float32 path is implemented as a first-class": (
            "The float32 path is implemented as a first-class memory-saving execution path. Float32 input is consumed directly by native CPU, Metal, or CUDA nearest-neighbour code when supported. KNN distances, graph or affinity weights, layouts, gradients, gains, updates, schedules, and FFT work arrays remain float32 through the computationally intensive stages. Metal PCA now copies float32 input directly into unified GPU storage, validates and centres/scales each feature with a native Metal reduction, performs MPS block-subspace TSVD products, and returns actual float32 scores and loadings. The final embedding is converted to an ordinary R numeric matrix only when required for compatibility. Total resident memory does not decrease by exactly 50% because integer indices, CSR row pointers, R object metadata, small projected eigensystems, and temporary workspaces are not double-precision payloads (Johnson et al., 2019; Douze et al., 2024; Caccia, 2026)."
        ),
        "The CUDA backend is implemented as a native fastEmbedR": (
            "The CUDA backend is implemented as a native fastEmbedR execution path rather than as a call to Python openTSNE, Python RAPIDS, or a relabelled CPU fallback. One-call CUDA KNN links directly to FAISS GPU exact search below 100,000 observations and to the RAPIDS cuVS C API for recall-tuned IVF-Flat above that threshold. The package-owned GPU KNN handle retains int32 indices and float32 distances on the selected device. The CUDA embedding bridge validates the non-self layout, metric, dimensions, and device before graph or affinity construction; unsupported layouts fail explicitly rather than being copied silently through R. Host KNN matrices and compatible externally created KNN objects remain supported by opentsne_knn() and umap_knn()."
        ),
        "For UMAP, the CUDA path consumes either host KNN matrices": (
            "For UMAP, the CUDA path consumes either host KNN matrices or a package-owned device-resident KNN handle. In the resident branch, device pointers are used to construct the graph, initialization, schedule, and optimizer state without materializing neighbour matrices in R. For graph_mode = \"fuzzy\", a CUDA kernel estimates the local connectivity offset and smooth KNN bandwidth for each row, computes directed membership strengths, and forms the symmetric fuzzy union. For graph_mode = \"binary\", the same neighbour topology is retained but valid symmetric graph weights are set to one before optimization. Valid entries are counted, prefix-scanned, and packed into compact COO arrays containing endpoints, float32 weights, and epochs_per_sample schedules."
        ),
        "The CUDA implementation performs explicit availability": (
            "The CUDA implementation performs explicit availability and memory checks before allocation. CUDA unavailability, unsupported KNN residency, unsupported distance storage, excessive neighbour width, and insufficient device memory are reported as errors rather than CPU fallbacks. The largest buffers are float32 arrays for distances, affinities, graph weights, layouts, gradients, updates, gains, schedules, and FFT grids; integer buffers hold neighbour indices and sparse graph offsets. Package-owned CUDA KNN buffers feed graph or affinity construction directly, while precomputed host KNN remains supported. FAISS, cuVS, cuFFT, RAFT, and CUDA remain system-level dependencies and their licenses are documented; fastEmbedR does not route its public native backends through Python (Johnson et al., 2019; Douze et al., 2024; Caccia, 2026)."
        ),
        "A recent implementation change was to make the source package": (
            "A recent implementation change was to make the source package suitable for Bioconductor-style submission without weakening accelerated workflows. The DESCRIPTION file treats faissR as an optional companion rather than a hard import, adds appropriate biocViews terms, and uses BiocStyle for vignettes. Runnable manual and vignette examples use small exact KNN objects or guarded optional calls, so the package can be checked without FAISS, cuVS, CUDA, or Metal. The current one-call CPU and Metal paths remain package-native; optional CUDA compilation links directly to the required system libraries. Explicit GPU requests still fail rather than silently relabelling CPU execution."
        ),
        "This study introduces fastEmbedR as a KNN-first embedding engine": (
            "This study introduces fastEmbedR as a KNN-first embedding engine for accelerated t-SNE and UMAP workflows in R. Its contribution is not limited to wrapping existing dimensionality-reduction software. From matrix input or a supplied KNN graph, fastEmbedR provides native KNN selection for one-call workflows, UMAP graph or openTSNE-style affinity construction, randomized-SVD initialization, CPU/Metal/CUDA layout optimization, backend metadata, and transformation workflows. CPU HNSW and Metal exact/IVF-Flat are implemented inside the package; CUDA builds call FAISS GPU and RAPIDS cuVS through native C++/CUDA interfaces, not Python. The optional faissR package remains a companion for reusable KNN objects. The t-SNE implementation follows the t-SNE, FIt-SNE, opt-SNE, and openTSNE literature at the level of objective, affinities, initialization, and scheduling, but the optimizer is implemented in fastEmbedR C++/Metal/CUDA (van der Maaten and Hinton, 2008; Chan et al., 2018; Linderman et al., 2019; Belkina et al., 2019b; Policar et al., 2024; Johnson et al., 2019; Douze et al., 2024; Caccia, 2026)."
        ),
        "fastEmbedR deliberately benefits from faissR": (
            "The total-runtime benchmark evaluates each package through its complete user-facing workflow. fastEmbedR one-call workflows include their package-native nearest-neighbour stage and native embedding stage; reference packages include their own internal search and optimization. The optional faissR interface is evaluated only when a reusable external KNN object is requested. KNN-entry functions remain valuable because they permit graph reuse across methods and make the embedding layer independently benchmarkable, but embedding-only timing is reported as a backend diagnostic rather than as the primary cross-package ranking."
        ),
        "The Metal backend requires an additional clarification.": (
            "The Metal backend now includes native exact and recall-tuned IVF-Flat nearest-neighbour search in addition to native PCA and embedding kernels. Thus, float32 matrix input can follow a package-native Apple GPU path without calling Python or the faissR R API. Metal uses unified memory, but avoiding explicit transfers does not remove command-dispatch, synchronization, FFT, and sparse-attraction costs. Total runtime and embedding-only runtime are therefore both reported: the former reflects user experience, while the latter isolates backend optimization from KNN and initialization."
        ),
        "The study has limitations.": (
            "The study has limitations. Runtime depends on hardware, memory bandwidth, feature dimension, KNN settings, initialization, and dataset structure, so speed-ups may differ across machines. Fixed seeds improve reproducibility, but atomic GPU updates need not be bitwise identical across architectures. CUDA builds require compatible FAISS, cuVS, RAFT, cuFFT, and CUDA system libraries; Metal acceleration requires Apple hardware and the Metal frameworks. Approximate native KNN routes are recall-tuned but remain data-dependent. Finally, implementation-level optimization diagnostics use cached KNN and initialization and therefore answer a different question from the primary total-runtime package comparison. These boundaries are reported explicitly so acceleration is not confused with silent fallback or reduced mathematical work (Amezquita et al., 2020; Heumos et al., 2023; Caccia, 2026)."
        ),
        "fastEmbedR provides a KNN-first and backend-transparent": (
            "fastEmbedR provides a KNN-first and backend-transparent solution for running t-SNE and UMAP workflows directly from R. Package-native CPU HNSW, Metal exact/IVF-Flat, graph and affinity construction, randomized-SVD initialization, and CPU/Metal/CUDA optimizers are exposed through a small API; CUDA builds link directly to FAISS, cuVS, RAFT, and cuFFT without Python. Reusable external KNN objects remain accepted through umap_knn() and opentsne_knn(). The controlled improvement strategy demonstrates that worker reuse, persistent plans and buffers, device-resident data flow, CUDA Graph capture, and native float32 PCA preprocessing can reduce runtime without reducing the algorithmic workload (van der Maaten and Hinton, 2008; Chan et al., 2018; Linderman et al., 2019; Belkina et al., 2019b; McInnes et al., 2018; Policar et al., 2024; Caccia, 2026)."
        ),
        "The package source is available at": (
            "The package source is available at https://github.com/tkcaccia/fastEmbedR. The source package is organized for Bioconductor-style checking: faissR is optional, biocViews and BiocStyle vignettes are provided, and examples can run without optional GPU libraries. Native one-call CPU and Metal KNN are included in fastEmbedR; optional CUDA builds link directly to system FAISS/cuVS/RAFT/CUDA libraries. The benchmark dataset inventory, command lines, session information, hardware metadata, random seeds, and accepted performance-engineering measurements are retained in the repository documentation and benchmark outputs. Total elapsed time is used for cross-package comparisons, whereas stage-level and cached-input measurements are explicitly labelled as backend diagnostics (Johnson et al., 2019; Douze et al., 2024; Caccia, 2026)."
        ),
    }
    for prefix, text in replacements.items():
        replace_paragraph(document, prefix, text)

    bioc_heading = find_paragraph(document, "2.3.2 Bioconductor-oriented")
    bioc_heading.text = "2.3.3 Bioconductor-oriented package boundary"
    improvement_heading = insert_paragraph_before(
        document,
        bioc_heading,
        "2.3.2 Performance engineering and change acceptance",
        "Heading 3",
    )
    methods_paragraphs = [
        "Optimization was performed as a sequence of controlled implementation changes rather than by reducing the mathematical workload. Each candidate retained the same observations, supplied KNN graph, metric, k or perplexity, initialization, random seed, t-SNE iteration schedule, UMAP epoch schedule, learning-rate and momentum rules, FFT-grid resolution, graph mode, and negative-sampling policy. KNN and PCA were cached when the embedding kernel was the object of study. This design separates improvements in memory layout, synchronization, data movement, and kernel scheduling from changes that could make an algorithm appear faster by performing less work.",
        "Candidates were accepted only after repeated wall-clock measurements and three quality gates. Deterministic CPU operations had to reproduce the accepted coordinates when operation order was unchanged. Parallel Metal and CUDA outputs had to remain within the variation of repeated atomic GPU runs, evaluated by coordinate correlation, Procrustes-aligned RMSD, neighbourhood overlap, and sampled distance correlation. Finally, label-coloured layouts from the complete dataset were inspected for cluster splitting, collapse, artificial holes, outliers, and loss of local organization. Slower candidates or candidates with weaker numerical or visual agreement were removed instead of becoming user options.",
        "CPU openTSNE now creates one reusable worker team per embedding, caches the FFT plan, retains per-worker scratch, and parallelizes independent grid-support operations while preserving stable point-splat accumulation. Metal keeps layout and optimizer buffers resident, encodes 16 unchanged iterations per command buffer, and performs float32 PCA validation, centering, scaling, and MPS TSVD products in unified GPU memory. CUDA retains cuFFT plans and workspaces and captures 25 unchanged optimizer iterations in CUDA Graph executables. These changes alter execution organization but not the UMAP or t-SNE objective, graph definition, force equations, or iteration count.",
        "Float32 was used for the large matrix products and persistent optimizer buffers because these stages are bandwidth- and capacity-sensitive. Integer indices and sparse offsets remain int32. Small metadata, convergence summaries, and projected eigensystems may use host numeric storage. For float32 Metal PCA input, scores and loadings are returned as true float32 matrices; conversion to R double is deferred to a final compatibility boundary only when required. Consequently, the largest floating-point payloads are halved, although total resident memory is not expected to fall by exactly 50%.",
    ]
    for text in methods_paragraphs:
        insert_paragraph_before(document, bioc_heading, text, "Body Text")
    improvement_heading.paragraph_format.keep_with_next = True

    discussion_heading = find_paragraph(document, "4. Discussion")
    results_heading = insert_paragraph_before(
        document,
        discussion_heading,
        "3.6 Implementation-level optimization validation",
        "Heading 2",
    )
    insert_paragraph_before(
        document,
        discussion_heading,
        "The controlled engineering benchmark used all 70,000 flattened MNIST observations, a cached KNN graph, a fixed PCA initialization, and unchanged public optimizer schedules. CPU used four threads unless the scaling row states otherwise. CPU/Metal and CUDA candidates were measured in independent paired development sessions; inference is based on the baseline-to-retained change within each row, not on absolute comparisons across rows or with Figure 3. These measurements isolate implementation changes and are not used to rank external packages.",
        "Body Text",
    )
    caption = insert_paragraph_before(
        document,
        discussion_heading,
        "Table 4. Retained implementation-only performance changes on MNIST70k. Baseline and retained runs used the same mathematical workload; negative percentages indicate lower elapsed time.",
        "Caption",
    )
    table = populate_improvement_table(document)
    discussion_heading._p.addprevious(table._tbl)
    insert_paragraph_before(
        document,
        discussion_heading,
        "The Metal PCA change moved the scan of the complete 70,000 by 784 float32 matrix from the host to a native Metal reduction. Warm median PCA time decreased from 0.429 to 0.104 seconds; the first invocation, including Metal pipeline compilation, required 0.774 seconds and is reported separately. The retained subspace had minimum canonical correlation 0.9980, maximum principal angle 3.65 degrees, and score-distance Pearson correlation 0.9977 against the CPU randomized-SVD reference. The 70,000 by 2 score object occupied approximately 0.56 MB rather than the approximately 1.12 MB double payload.",
        "Body Text",
    )
    insert_paragraph_before(
        document,
        discussion_heading,
        "Plausible but unsuccessful changes were removed. Private per-thread CPU splat grids increased full-MNIST median openTSNE time to 14.897 seconds; fused Metal post-update centering increased it to 2.613 seconds; and moving the small PCA orthonormalization into a separate Metal kernel increased median PCA time to 0.449 seconds. Reporting these rejections is important because fewer kernel launches or more parallel code do not necessarily improve end-to-end performance.",
        "Body Text",
    )
    results_heading.paragraph_format.keep_with_next = True
    caption.paragraph_format.keep_with_next = True

    metal_discussion = find_paragraph(document, "The Metal backend now includes")
    insert_paragraph_before(
        document,
        metal_discussion,
        "The improvement results also show why backend optimization must be profile-guided. CPU gains came primarily from worker, plan, and scratch reuse and scaled strongly through four performance cores. Metal PCA gained substantially by eliminating a host scan of the full input, whereas the already resident Metal openTSNE loop gained only modestly from command batching. CUDA benefited most from CUDA Graph launch amortization. Candidate fusion and parallelization experiments that increased memory traffic were slower despite appearing more GPU-oriented. The retained strategy therefore optimizes the measured bottleneck of each backend rather than imposing one implementation pattern on all devices.",
        "Normal",
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    update_document(args.source.expanduser(), args.output.expanduser())


if __name__ == "__main__":
    main()
