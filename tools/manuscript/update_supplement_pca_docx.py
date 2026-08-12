#!/usr/bin/env python3
"""Synchronize the PCA section of the MLOSS supplement DOCX."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def find_paragraph(document: Document, prefix: str):
    normalized_prefix = prefix.replace("\u00a0", " ")
    for paragraph in document.paragraphs:
        if paragraph.text.replace("\u00a0", " ").startswith(normalized_prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def format_table(table) -> None:
    table.style = "Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.52), Inches(0.72), Inches(0.72), Inches(1.62), Inches(1.14)]
    for row_index, row in enumerate(table.rows):
        keep_row_together(row)
        for cell, width in zip(row.cells, widths):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True
    set_repeat_table_header(table.rows[0])


def update_table_caption_numbers(document: Document) -> None:
    pattern = re.compile(r"^Table S(\d+)\.")
    for paragraph in document.paragraphs:
        match = pattern.match(paragraph.text)
        if match and int(match.group(1)) >= 5:
            number = int(match.group(1)) + 1
            set_text(paragraph, pattern.sub(f"Table S{number}.", paragraph.text, count=1))


def main() -> int:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: update_supplement_pca_docx.py PATH_TO_DOCX")

    path = Path(sys.argv[1]).resolve()
    document = Document(path)

    has_backend_speed_table = any(
        paragraph.text.startswith("Table S5. MNIST70k PCA runtime")
        for paragraph in document.paragraphs
    )
    if not has_backend_speed_table:
        update_table_caption_numbers(document)

    paragraph = find_paragraph(document, "Table\u00a03 distinguishes")
    set_text(
        paragraph,
        'Table 3 distinguishes the three execution routes. Public CPU and CUDA '
        'calls use a compatible optional RSVD provider when one is installed and '
        'otherwise use the package-native RSVD fallback. For backend="cpu", the '
        'exported n.cores argument sets the temporary BLAS/OpenMP core limit and '
        'the fit records the requested and observable counts; a single-threaded '
        'BLAS remains limited to one core. The package-native CUDA fallback '
        'accelerates the large matrix products but completes QR and the small SVD '
        'on the host; by contrast, the one-call CUDA openTSNE initializer calls '
        'RAFT TSVD directly and retains its two-dimensional score buffer on-device. '
        'Metal uses a separate package-native float32 block-subspace TSVD and MPS '
        'matrix products with unified GPU buffers. No route calls Python, and an '
        'unavailable requested GPU backend raises an error rather than being '
        'reported as GPU work. irlba is not a package dependency or implementation '
        'route; it is used only as an external benchmark reference.',
    )

    paragraph = find_paragraph(document, "The CPU route is therefore multicore-capable")
    set_text(
        paragraph,
        'The CPU route is therefore multicore-capable, but the effective count '
        'still depends on the R process and linked BLAS/LAPACK. The public n.cores '
        'argument temporarily sets OMP_NUM_THREADS, OPENBLAS_NUM_THREADS, '
        'MKL_NUM_THREADS, and related limits and uses RhpcBLASctl when installed; '
        'previous settings are restored when the call ends. The fit reports '
        'n.cores_requested, n.cores_effective, and core_control. A build linked to '
        'a single-threaded BLAS remains single-threaded, which is reported rather '
        'than hidden.',
    )

    paragraph = find_paragraph(document, "The production KNN and embedding paths")
    set_text(
        paragraph,
        'The production KNN and embedding paths use float32 numerical buffers. For '
        'ordinary R matrices, a controlled conversion is required at the native '
        'boundary; for float::float32 input, the package reads the float payload '
        'directly. KNN distances, graph or affinity weights, layouts, gradients, '
        'gains, and update buffers remain float32, while indices and CSR offsets '
        'use int32. Backend-specific PCA may use an optional RSVD provider or '
        'host-side setup before entering its float32 production path, so the '
        'precision claim does not extend to every diagnostic helper or temporary '
        'object. If the input is float32, fastEmbedR returns a float32 embedding; '
        'if the input is a standard R matrix, only the final layout is converted '
        'back to double.',
    )

    paragraph = find_paragraph(document, "The validated numerical package tree")
    set_text(
        paragraph,
        'The validated numerical package tree is frozen at fastEmbedR commit '
        '58e39d53d38. The validated multi-architecture CUDA image and embedded '
        'shared library have SHA-256 prefixes 8a91df71be0d and 7975ae24f952, '
        'respectively. Image labels also pin the nearest-neighbor and PCA companion '
        'libraries, KODAMA C++, and KODAMA R. Full, unabridged commits and hashes '
        'are archived in generated/release_identity.txt. Runtime smoke tests '
        'covered CPU/CUDA float32 KNN, PCA, openTSNE, UMAP, KODAMA, Python reference '
        'methods, and FIt-SNE. Later repository commits add manuscript and '
        'generated publication artifacts without changing the validated numerical '
        'package tree.',
    )

    for table in document.tables:
        for row in table.rows:
            if len(row.cells) == 3 and row.cells[0].text == "PCA initialization":
                row.cells[1].text = (
                    "Centering/scaling, randomized singular value decomposition "
                    "(RSVD), Metal block-subspace truncated singular value "
                    "decomposition (TSVD), and t-SNE rescaling"
                )
                row.cells[2].text = (
                    "Compatible optional RSVD provider for public CPU/CUDA PCA, with "
                    "a package-native fallback; Metal Performance Shaders (MPS) "
                    "matrix products; RAFT TSVD for the resident CUDA openTSNE "
                    "initializer"
                )

    backend_table = next(
        table
        for table in document.tables
        if table.cell(0, 0).text == "Backend"
        and table.cell(0, 1).text == "Production implementation"
    )
    backend_table.cell(1, 1).text = (
        "Compatible optional RSVD provider or package-native Gaussian-sketch RSVD"
    )
    backend_table.cell(3, 1).text = (
        "Compatible optional CUDA RSVD provider for public pca(); RAFT TSVD for "
        "resident openTSNE initialization"
    )

    for paragraph in document.paragraphs:
        replacements = {
            "Tables\u00a07\u20139": "Tables 8\u201310",
            "Table\u00a010": "Table 11",
            "Table\u00a011": "Table 12",
            "Table\u00a013": "Table 14",
            "Table\u00a06": "Table 7",
            "Table\u00a017": "Table 18",
        }
        text = paragraph.text
        revised = text
        for old, new in replacements.items():
            revised = revised.replace(old, new)
        if revised != text:
            set_text(paragraph, revised)

    if has_backend_speed_table:
        target = find_paragraph(document, "On the Apple M3 system")
    else:
        target = find_paragraph(document, "A separate Mac benchmark")
        caption = target.insert_paragraph_before(
            "Table S5. MNIST70k PCA runtime by fastEmbedR backend. Values are "
            "medians over seeds 4, 17, and 42; parentheses give IQRs. Speedup is "
            "relative to the one-core CPU route on the same machine. The Mac and "
            "HPC blocks are not compared with each other.",
            style="Table Caption",
        )
        caption.paragraph_format.keep_with_next = True

        rows = [
            ["Machine", "Backend", "CPU cores", "Runtime, s", "Speedup vs. CPU-1"],
            ["MacBook Pro (M3)", "CPU", "1", "1.781 (1.742–1.853)", "1.00"],
            ["MacBook Pro (M3)", "CPU", "4", "1.699 (1.667–1.715)", "1.05"],
            ["MacBook Pro (M3)", "Metal", "–", "0.274 (0.255–0.345)", "6.50"],
            ["HPC Xeon/L40S", "CPU", "1", "2.404 (2.365–2.406)", "1.00"],
            ["HPC Xeon/L40S", "CPU", "4", "2.096 (2.093–2.097)", "1.15"],
            ["HPC Xeon/L40S", "CUDA", "–", "2.794 (2.780–2.795)", "0.86"],
        ]
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        for row, values in zip(table.rows, rows):
            for cell, value in zip(row.cells, values):
                cell.text = value
        format_table(table)
        target._p.addprevious(table._tbl)

    set_text(
        target,
        'On the Apple M3 system, Metal reduced the isolated PCA call from 1.781 '
        'seconds at one CPU core to 0.274 seconds, a within-machine speedup of 6.50. '
        'On the Xeon/L40S system, the public CUDA pca() call required 2.794 seconds '
        'compared with 2.404 and 2.096 seconds for one and four CPU cores. Thus '
        'backend-native PCA is not uniformly faster: the public CUDA route pays '
        'setup and transfer costs for this matrix shape, while its principal '
        'advantage in complete CUDA openTSNE is retaining the initialization '
        'on-device. Median Procrustes correlations against irlba were 0.9978 for '
        'the local CPU route and 0.9991 for Metal. These results demonstrate that '
        'PCA speed depends on matrix shape, numerical libraries, compiler, and '
        'hardware; the contribution is a backend-consistent PCA and openTSNE '
        'initialization boundary rather than a universal speed claim.',
    )

    for paragraph in document.paragraphs:
        if "fastPLS" in paragraph.text:
            raise RuntimeError(f"Unremoved fastPLS reference: {paragraph.text}")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if "fastPLS" in cell.text:
                    raise RuntimeError(f"Unremoved fastPLS reference: {cell.text}")

    temporary = path.with_suffix(".tmp.docx")
    document.save(temporary)
    temporary.replace(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
