#!/usr/bin/env python3
"""Synchronize submission DOCX files with the final LaTeX text and figures."""

from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
MLOSS = ROOT / "manuscript" / "mloss"
DELIVERABLES = MLOSS / "deliverables"


def find_paragraph(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def drawing_target(document: Document, paragraph) -> str:
    blips = paragraph._p.xpath(".//a:blip")
    if len(blips) != 1:
        raise RuntimeError("Expected exactly one drawing in figure paragraph")
    relationship_id = blips[0].get(qn("r:embed"))
    return "word/" + document.part.rels[relationship_id].target_ref


def replace_zip_members(path: Path, replacements: dict[str, Path]) -> None:
    with tempfile.TemporaryDirectory(prefix="fastembedr-docx-") as temporary:
        rewritten = Path(temporary) / path.name
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
            rewritten, "w", zipfile.ZIP_DEFLATED
        ) as destination:
            for item in source.infolist():
                if item.filename not in replacements:
                    destination.writestr(item, source.read(item.filename))
            for member, replacement in replacements.items():
                destination.write(replacement, member)
        shutil.copy2(rewritten, path)


def update_main_document() -> None:
    path = DELIVERABLES / "fastEmbedR_MLOSS_manuscript.docx"
    document = Document(path)

    find_paragraph(document, "For matrix input, CPU nearest-neighbor search").text = (
        "For matrix input, CPU nearest-neighbor search uses package-resident "
        "hierarchical navigable small world (HNSW), distilled from permissively "
        "licensed FAISS code (Johnson, Douze, and Jégou 2021; Douze et al. 2024); "
        "Metal uses exact or recall-tuned inverted-file flat (IVF-Flat); CUDA "
        "links FAISS GPU exact or cuVS IVF-Flat search and can retain KNN "
        "on-device. The exported pca() uses randomized singular value "
        "decomposition (RSVD) (Halko, Martinsson, and Tropp 2011): CPU and CUDA "
        "use package-owned RSVD orchestration, Metal uses float32 block-subspace "
        "truncated singular value decomposition (TSVD) with Metal Performance "
        "Shaders (MPS), and resident CUDA openTSNE initialization uses RAFT "
        "TSVD. No PCA route delegates to another R package. Setting "
        "opentsne_init=TRUE produces the small-scale PCA initialization used by "
        "t-SNE (Kobak and Berens 2019)."
    )

    first_figure = find_paragraph(document, "Figure 1.")
    second_figure = find_paragraph(document, "Figure 2.")
    first_figure.text = (
        "Figure 1. Runtime for the principal t-SNE comparisons. The single "
        "panel includes CPU and CUDA results, identified explicitly in the "
        "legend, for Rtsne, FIt-SNE, fastEmbedR openTSNE, direct Python "
        "openTSNE, and direct RAPIDS cuML t-SNE. Gray denotes established R "
        "references, blue denotes fastEmbedR, purple denotes Python openTSNE, "
        "and orange denotes RAPIDS. The continuous pseudo-log runtime axis "
        "retains a true zero baseline. Bars are medians over seeds 4, 17, and "
        "42; error bars span the interquartile range. R rows report total "
        "public-function time, whereas direct-Python rows report Python fit "
        "time only; these boundaries are not pooled. The exhaustive figure, "
        "including R-mediated Python timings, is supplied in the supplement. "
        "FlowRepository has no successful Rtsne row, so its unmatched "
        "openTSNE result is descriptive only and excluded from the paired "
        "Rtsne comparison."
    )
    second_figure.text = (
        "Figure 2. Runtime for the principal UMAP comparisons. The single "
        "panel combines CPU and CUDA results, identified in the legend, for "
        "uwot fast SGD, fastEmbedR fuzzy UMAP, direct Python umap-learn, and "
        "direct RAPIDS cuML UMAP. Colors follow Figure 1. R rows report total "
        "public-function time, whereas direct-Python rows report Python fit "
        "time only; these scopes remain separate. The pseudo-log runtime axis "
        "retains a true zero baseline. Exhaustive comparisons with the R umap "
        "package, uwot default, binary fastEmbedR UMAP, and R-mediated Python "
        "timings are supplied in the supplement. FlowRepository CPU UMAP was "
        "not reached before the enclosing job ended, so its CUDA-only "
        "fastEmbedR rows are descriptive rather than paired CPU-reference "
        "comparisons."
    )

    references = find_paragraph(document, "References")
    previous = references._p.getprevious()
    acknowledgment_text = (
        "The authors thank the University of Cape Town's ICTS High Performance "
        "Computing team for providing a high-performance computing facility for "
        "this study (https://ucthpc.uct.ac.za/)."
    )
    if previous is None or "University of Cape Town's ICTS" not in "".join(
        previous.itertext()
    ):
        paragraph = references.insert_paragraph_before(acknowledgment_text)
        heading = references.insert_paragraph_before(
            "Acknowledgments and Disclosure of Funding"
        )
        heading.style = references.style

    drawing_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.xpath(".//a:blip")
    ]
    if len(drawing_paragraphs) < 2:
        raise RuntimeError("Main DOCX does not contain both runtime figures")
    first_target = drawing_target(document, drawing_paragraphs[0])
    second_target = drawing_target(document, drawing_paragraphs[1])
    document.save(path)
    replace_zip_members(
        path,
        {
            first_target: MLOSS / "figures" / "runtime_tsne_main_methods.png",
            second_target: MLOSS / "figures" / "runtime_umap_main_methods.png",
        },
    )


def update_supplement_document() -> None:
    path = DELIVERABLES / "fastEmbedR_MLOSS_supplement.docx"
    document = Document(path)

    for paragraph in document.paragraphs:
        match = re.match(r"Figure S(\d+)\.", paragraph.text)
        if match:
            old_number = int(match.group(1))
            paragraph.text = re.sub(
                r"^Figure S\d+\.",
                f"Figure S{old_number + 2}.",
                paragraph.text,
                count=1,
            )

    landmark_text = find_paragraph(document, "The landmark benchmark used")
    landmark_text.text = landmark_text.text.replace(
        "Table S8 and Figure 1", "Table S8 and Figure S3"
    )

    python_summary = find_paragraph(document, "The Python benchmark branch")
    python_summary.text = (
        "The Python benchmark branch evaluated Python openTSNE, umap-learn, "
        "and RAPIDS cuML t-SNE/UMAP both through reticulate and, where "
        "available, as a direct Python subprocess. Timing modes remain "
        "separate: native Python fit time is not presented as equivalent to an "
        "R-mediated total call. Main-text Figures 1 and 2 use a reduced "
        "principal-method set for readability. Supplementary Figures S1 and "
        "S2 retain every successful t-SNE and UMAP implementation, "
        "respectively, combining CPU and CUDA in one panel while retaining "
        "backend labels and distinct direct-Python and R-mediated entries. "
        "The machine-readable files identify each row with timing_scope and "
        "runtime_measure. Values remain in separate "
        "r_mediated_total_call_sec, direct_python_fit_sec, and "
        "direct_python_process_total_sec columns. The complete three-seed "
        "Python summary is archived as generated/python_summary_median.csv. "
        "The exhaustive rows are archived in "
        "generated/runtime_tsne_all_methods.csv and "
        "generated/runtime_umap_all_methods.csv; main-figure rows are in the "
        "corresponding runtime_*_main_methods.csv files."
    )

    target = find_paragraph(
        document, "S12 Reference-Implementation and Backend Validation"
    )
    heading = target.insert_paragraph_before("Exhaustive Runtime Comparisons")
    heading.style = target.style

    tsne_plot = target.insert_paragraph_before()
    tsne_plot.alignment = 1
    tsne_plot.add_run().add_picture(
        str(MLOSS / "figures" / "runtime_tsne_all_methods.png"),
        width=Inches(6.4),
    )
    tsne_caption = target.insert_paragraph_before(
        "Figure S1. Exhaustive t-SNE runtime comparison. CPU and CUDA results "
        "share one panel. Bars are medians over seeds 4, 17, and 42 and error "
        "bars span the interquartile range. R-mediated total-call and direct "
        "Python fit timings are labeled separately; missing rows remain empty."
    )
    tsne_caption.style = find_paragraph(document, "Figure S3.").style

    umap_plot = target.insert_paragraph_before()
    umap_plot.alignment = 1
    umap_plot.add_run().add_picture(
        str(MLOSS / "figures" / "runtime_umap_all_methods.png"),
        width=Inches(6.4),
    )
    umap_caption = target.insert_paragraph_before(
        "Figure S2. Exhaustive UMAP runtime comparison. CPU and CUDA results "
        "share one panel. Bars are medians over seeds 4, 17, and 42 and error "
        "bars span the interquartile range. R-mediated total-call and direct "
        "Python fit timings are labeled separately; missing rows remain empty."
    )
    umap_caption.style = find_paragraph(document, "Figure S3.").style

    document.save(path)


def main() -> None:
    update_main_document()
    update_supplement_document()


if __name__ == "__main__":
    main()
